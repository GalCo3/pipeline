"""Demo producer: uploads sample documents to S3 (MinIO) and produces matching
CargoMessage records to Kafka, so the cargo-lexical service has something to consume.

Every message mirrors `services/cargo-lexical/config/models.py::CargoMessage`, including
its strict `%Y-%m-%dT%H:%M:%S` datetime format.
"""

from __future__ import annotations

import io
import json
import logging
import os
import socket
import time
import zlib
from collections.abc import Callable
from datetime import UTC, datetime

import boto3
from botocore.exceptions import ClientError
from confluent_kafka import Producer

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("demo-producer")

BOOTSTRAP_SERVERS = os.environ.get("BOOTSTRAP_SERVERS", "kafka:9092")
TOPIC = os.environ.get("TOPIC", "cargo-lexical.files")
S3_ENDPOINT = os.environ.get("S3_ENDPOINT", "http://minio:9000")
S3_ACCESS_KEY = os.environ.get("S3_ACCESS_KEY", "minioadmin")
S3_SECRET_KEY = os.environ.get("S3_SECRET_KEY", "minioadmin")
S3_BUCKET = os.environ.get("S3_BUCKET", "cargo-lexical")
# CargoMessage parses every timestamp with this exact format — no timezone,
# no microseconds — so isoformat() output would be rejected.
DATETIME_FORMAT = "%Y-%m-%dT%H:%M:%S"
# 0 = produce the demo batch once and exit; >0 = keep re-producing every N seconds.
INTERVAL_SECONDS = float(os.environ.get("INTERVAL_SECONDS", "0"))
# Keeps document ids and S3 keys distinct across re-runs of the job, so a rerun
# adds documents instead of overwriting the previous batch.
RUN_SEED = int(os.environ.get("RUN_SEED", "") or zlib.crc32(socket.gethostname().encode()) % 90_000)


def build_txt() -> bytes:
    return (
        b"Cargo demo plain-text document.\n"
        b"Extracted by the streaming text extractor (no Tika needed).\n"
        b"Line three, with some searchable words: manifest, container, shipment.\n"
    )


def build_docx() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Cargo Demo DOCX", level=1)
    document.add_paragraph("Shipment manifest for container HERMES-0001.")
    document.add_paragraph("Handled by the buffered DOCX extractor.")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf() -> bytes:
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 800, "Cargo Demo PDF")
    pdf.drawString(72, 780, "Bill of lading for container HERMES-0002.")
    pdf.save()
    return buffer.getvalue()


def build_xlsx() -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "manifest"
    sheet.append(["container", "weight_kg", "destination"])
    sheet.append(["HERMES-0003", 1200, "Haifa"])
    sheet.append(["HERMES-0004", 890, "Ashdod"])

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


OOXML = "application/vnd.openxmlformats-officedocument"

DEMO_FILES: list[tuple[str, str, Callable[[], bytes]]] = [
    ("manifest.txt", "text/plain", build_txt),
    ("shipment.docx", f"{OOXML}.wordprocessingml.document", build_docx),
    ("bill-of-lading.pdf", "application/pdf", build_pdf),
    ("weights.xlsx", f"{OOXML}.spreadsheetml.sheet", build_xlsx),
]


def create_s3_client():
    return boto3.client(
        "s3",
        aws_access_key_id=S3_ACCESS_KEY,
        aws_secret_access_key=S3_SECRET_KEY,
        endpoint_url=S3_ENDPOINT,
    )


def ensure_bucket(s3) -> None:
    try:
        s3.head_bucket(Bucket=S3_BUCKET)
    except ClientError:
        s3.create_bucket(Bucket=S3_BUCKET)
        logger.info("Created bucket %s", S3_BUCKET)


def upload_demo_files(s3, run_id: int) -> list[tuple[str, str]]:
    """Uploads the demo files under a per-run prefix. Returns (s3_key, name) pairs."""
    uploaded = []
    for name, content_type, build in DEMO_FILES:
        body = build()
        key = f"demo/run-{RUN_SEED}-{run_id}/{name}"
        s3.put_object(Bucket=S3_BUCKET, Key=key, Body=body, ContentType=content_type)
        logger.info("Uploaded s3://%s/%s (%d bytes)", S3_BUCKET, key, len(body))
        uploaded.append((key, name))
    return uploaded


def build_message(doc_id: int, s3_key: str, name: str) -> dict:
    now = datetime.now(UTC).strftime(DATETIME_FORMAT)
    path_id = f"1/42/{doc_id}"

    return {
        "id": doc_id,
        "name": name,
        "holder": "demo-holder",
        "description": f"Demo cargo-lexical document {name}",
        "is_verified": True,
        "file_labels": [
            {
                "label_id": 1,
                "label_name": "demo",
                "created": now,
                "group_id": 10,
                "group_name": "demo-group",
            }
        ],
        "path_id": path_id,
        "path": "/root/demo/cargo-lexical",
        "reality_id": "reality-1",
        "reality_type": "operational",
        "s3_key": s3_key,
        "s3_bucket": S3_BUCKET,
        "created": now,
        "last_modified": now,
        "ver_last_modified": now,
        # A non-null delete_date tells the service to delete the document, and
        # last_modified > ver_last_modified makes it an update instead of an
        # enrichment — keep both neutral so the demo exercises the index path.
        "delete_date": None,
    }


def delivery_report(err, msg) -> None:
    if err is not None:
        logger.error("Delivery failed for key %s: %s", msg.key(), err)
    else:
        logger.info("Produced to %s [%s] @ %s", msg.topic(), msg.partition(), msg.offset())


def main() -> None:
    s3 = create_s3_client()
    ensure_bucket(s3)

    producer = Producer({"bootstrap.servers": BOOTSTRAP_SERVERS})

    run_id = 0
    while True:
        run_id += 1
        for index, (key, name) in enumerate(upload_demo_files(s3, run_id), start=1):
            doc_id = RUN_SEED * 1_000 + run_id * 100 + index
            message = build_message(doc_id, key, name)
            producer.produce(
                TOPIC,
                key=str(doc_id),
                value=json.dumps(message).encode("utf-8"),
                on_delivery=delivery_report,
            )
        producer.flush(10)
        logger.info("Batch %d produced to topic %s", run_id, TOPIC)

        if INTERVAL_SECONDS <= 0:
            return
        time.sleep(INTERVAL_SECONDS)


if __name__ == "__main__":
    main()
