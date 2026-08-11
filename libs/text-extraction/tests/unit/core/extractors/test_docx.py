import io

import pytest
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from tests.fakes import (
    FakeDocument,
    FakeHeaderFooter,
    FakeParagraph,
    FakeSection,
    FakeTable,
)
from tests.mocks import ExtractorMockConfigurator

from hermes.text_extraction.core.extractors import docx
from hermes.text_extraction.exceptions import CorruptDocumentError, TextLimitExceededError


@pytest.mark.parametrize(
    "fake_doc, expected",
    [
        # 1. Simple text (just paragraphs, no tables, no header/footer customized)
        (
            FakeDocument(
                content=[FakeParagraph("Stubbed Paragraph 1"), FakeParagraph("Stubbed Paragraph 2")]
            ),
            "Stubbed Paragraph 1\nStubbed Paragraph 2",
        ),
        # 2. Tables only
        (
            FakeDocument(content=[FakeTable([["Cell 1A", "Cell 1B"], ["Cell 2A", "Cell 2B"]])]),
            "Cell 1A | Cell 1B\nCell 2A | Cell 2B",
        ),
        # 3. Paragraphs and tables (intermediate combination)
        (
            FakeDocument(
                content=[FakeParagraph("Hello Paragraph Text"), FakeTable([["Cell1", "Cell2"]])]
            ),
            "Hello Paragraph Text\nCell1 | Cell2",
        ),
        # 4. Complex document (headers + footers + paragraphs + tables)
        (
            FakeDocument(
                content=[FakeParagraph("Body Text Content"), FakeTable([["Cell1", "Cell2"]])],
                sections=[
                    FakeSection(
                        header=FakeHeaderFooter(content=[FakeParagraph("Header Text Content")]),
                        footer=FakeHeaderFooter(content=[FakeParagraph("Footer Text Content")]),
                    )
                ],
            ),
            "Header Text Content\nFooter Text Content\nBody Text Content\nCell1 | Cell2",
        ),
    ],
    ids=[
        "paragraphs_only",
        "tables_only",
        "paragraphs_and_tables",
        "complex_document",
    ],
)
def test_docx_extract_success(
    mock_docx: ExtractorMockConfigurator,
    fake_doc: FakeDocument,
    expected: str,
) -> None:
    # Arrange
    mock_docx(document=fake_doc)
    payload = io.BytesIO(b"docx-dummy-data")

    # Act
    result = docx.extract(payload, 100)

    # Assert
    assert result == expected


def test_docx_extract_corrupt(mock_docx: ExtractorMockConfigurator) -> None:
    # Arrange
    mock_docx(side_effect=PackageNotFoundError("Corrupted package structure"))
    payload = io.BytesIO(b"corrupt-docx-data")

    # Act & Assert
    with pytest.raises(CorruptDocumentError) as exc_info:
        docx.extract(payload, 100)

    assert (
        exc_info.value.mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_docx_extract_limit_exceeded(mock_docx: ExtractorMockConfigurator) -> None:
    # Arrange
    fake_doc = FakeDocument(
        content=[
            FakeParagraph("Hello Paragraph Text"),
            FakeTable([["Cell1", "Cell2"]]),
        ]
    )
    mock_docx(document=fake_doc)
    payload = io.BytesIO(b"docx-dummy-data")

    # Act & Assert
    with pytest.raises(TextLimitExceededError):
        docx.extract(payload, 10)


# =====================================================================
# Internal Helper Unit Tests
# =====================================================================


def test_iter_block_items() -> None:
    # Arrange
    # 1. Section with header/footer
    header = FakeHeaderFooter(content=[FakeParagraph("Header Text")])
    footer = FakeHeaderFooter(content=[FakeParagraph("Footer Text")])
    section = FakeSection(header=header, footer=footer)

    # 2. Body items
    p = FakeParagraph("Body Text")
    table = FakeTable([["Cell Text"]])

    doc = FakeDocument(content=[p, table], sections=[section])

    # Act
    items = list(docx._iter_block_items(doc))

    # Assert
    assert len(items) == 4
    assert isinstance(items[0], Paragraph)
    assert items[0].text == "Header Text"
    assert isinstance(items[1], Paragraph)
    assert items[1].text == "Footer Text"
    assert isinstance(items[2], Paragraph)
    assert items[2].text == "Body Text"
    assert isinstance(items[3], Table)
    # Check rows cell text
    assert items[3].rows[0].cells[0].text == "Cell Text"


def test_extract_item_text_paragraph() -> None:
    # Arrange
    p = FakeParagraph("   Some paragraph text   ")

    # Act
    chunks = list(docx._extract_item_text(p))

    # Assert
    assert chunks == ["Some paragraph text"]


def test_extract_item_text_paragraph_empty() -> None:
    # Arrange
    p = FakeParagraph("")

    # Act
    chunks = list(docx._extract_item_text(p))

    # Assert
    assert chunks == []


def test_extract_item_text_table() -> None:
    # Arrange
    table = FakeTable([["  Cell 1A  ", ""], ["Cell 2A", "Cell 2B"]])

    # Act
    chunks = list(docx._extract_item_text(table))

    # Assert
    assert chunks == ["Cell 1A", "Cell 2A | Cell 2B"]
